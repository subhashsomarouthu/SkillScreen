import logging
from typing import Optional, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
import io

logger = logging.getLogger(__name__)

class TextExtractor:
    """Extracts text content from various file formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_docx  # Will try to handle .doc as .docx
        }
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from file based on its extension"""
        file_path_obj = Path(file_path)
        file_ext = file_path_obj.suffix.lower()
        
        if file_ext not in self.supported_formats:
            return {
                "success": False,
                "error": f"Unsupported file format: {file_ext}",
                "text": ""
            }
        
        try:
            extractor_func = self.supported_formats[file_ext]
            text = extractor_func(file_path)
            
            return {
                "success": True,
                "text": text,
                "char_count": len(text)
            }
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise
        
        return text.strip()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            raise
        
        return text.strip()
    
    def is_text_extractable(self, file_path: str) -> bool:
        """Check if file can have text extracted"""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_formats
